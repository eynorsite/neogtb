#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Génère le lead magnet .docx du modèle de CCTP GTB vendor-neutral à partir des
sources Markdown de docs/cctp/ (CCTP + grille de dépouillement).

Source de vérité = docs/cctp/*.md (versionné). Ce script ne fait que la mise en
forme .docx ; aucun contenu réglementaire n'est ajouté ici.

Dépendances : python-docx, markdown, htmldocx (pip install --user).
"""

import os
import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from htmldocx import HtmlToDocx

ROOT = "/Users/calmoulrich/ProjetsWeb/neogtb"
SRC_CCTP = os.path.join(ROOT, "docs/cctp/CCTP-GTB-multiprotocole-vendor-neutral.md")
SRC_GRILLE = os.path.join(ROOT, "docs/cctp/grille-depouillement-GTB.md")
OUT = os.path.join(ROOT, "admin/resources/lead-magnets/modele-cctp-decret-bacs-neogtb.docx")

DARK = RGBColor(0x0F, 0x17, 0x2A)
ACCENT = RGBColor(0x26, 0x7A, 0x43)

MD_EXT = ["tables", "fenced_code", "sane_lists"]


def md_to_html(path):
    with open(path, encoding="utf-8") as f:
        return markdown.markdown(f.read(), extensions=MD_EXT)


doc = Document()

# Police de base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

# Page de garde minimale
t = doc.add_heading("Modèle de CCTP — GTB multi-protocole (vendor-neutral)", level=0)
for r in t.runs:
    r.font.color.rgb = DARK
p = doc.add_paragraph()
r = p.add_run("Cahier des Clauses Techniques Particulières + grille de dépouillement — trame type réutilisable")
r.italic = True
r.font.color.rgb = ACCENT
p2 = doc.add_paragraph()
r2 = p2.add_run(
    "Ressource NeoGTB — neogtb.fr. Document neutre, n'avantageant aucune marque ni protocole. "
    "À adapter projet par projet : renseigner les champs [À COMPLÉTER], traiter les blocs OPTION et "
    "la check-list de l'article 13 (points réglementaires à reconfirmer sur Légifrance/AFNOR à la date de diffusion)."
)
r2.font.size = Pt(9)
doc.add_page_break()

parser = HtmlToDocx()
parser.add_html_to_document(md_to_html(SRC_CCTP), doc)
doc.add_page_break()
parser.add_html_to_document(md_to_html(SRC_GRILLE), doc)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)

# Stats de contrôle
tables = len(doc.tables)
paras = len(doc.paragraphs)
print(f"OK -> {OUT}")
print(f"paragraphes={paras} tableaux={tables}")
