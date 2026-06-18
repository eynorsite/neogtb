#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Génère le modèle de CCTP décret BACS (classe B) au format .docx — NeoGTB."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
import os

OUT = "/Users/calmoulrich/ProjetsWeb/neogtb/admin/resources/lead-magnets/modele-cctp-decret-bacs-neogtb.docx"

ACCENT = RGBColor(0x26, 0x7a, 0x43)   # accent-600 NeoGTB
DARK = RGBColor(0x0f, 0x17, 0x2a)

doc = Document()

# Styles de base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x33, 0x3a, 0x45)

def todo(p, text):
    """Ajoute un placeholder surligné [à compléter]."""
    run = p.add_run(text)
    run.bold = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return run

def para(text="", bold=False, italic=False, size=11, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p

def numbered(text):
    p = doc.add_paragraph(text, style="List Number")
    return p

def article(num, title):
    p = doc.add_paragraph()
    r = p.add_run("ARTICLE %d" % num)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_after = Pt(0)
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = DARK

# ---------- PAGE DE GARDE ----------
title = doc.add_heading("Modèle de CCTP — Système de GTB conforme au décret BACS", level=0)
for run in title.runs:
    run.font.color.rgb = DARK

para("Cahier des Clauses Techniques Particulières (CCTP) — trame type",
     bold=True, size=13, color=ACCENT, space_after=2)
para("Système d'automatisation et de contrôle du bâtiment (BACS) visant la classe B "
     "au sens de la norme NF EN ISO 52120-1.", italic=True, size=11, space_after=12)

p = doc.add_paragraph()
todo(p, "[À COMPLÉTER : désignation du bâtiment, maître d'ouvrage, adresse, numéro de consultation]")
p.paragraph_format.space_after = Pt(12)

para("Ressource gratuite NeoGTB — neogtb.fr — Conseil GTB indépendant, sans lien fabricant.",
     size=9, italic=True, space_after=12)

# Avertissement
warn = doc.add_paragraph()
wr = warn.add_run("Avertissement — ")
wr.bold = True
warn.add_run("Ce document est une trame indicative à adapter à chaque projet. Il ne constitue ni "
             "un avis juridique, ni une garantie de conformité. Les passages surlignés en jaune "
             "doivent être renseignés selon le bâtiment concerné ; supprimez ce qui ne s'applique pas.")
warn.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ---------- SOMMAIRE ----------
h = doc.add_heading("Sommaire", level=1)
for run in h.runs:
    run.font.color.rgb = DARK
sommaire = [
    "Article 1 — Objet du marché",
    "Article 2 — Cadre réglementaire et normatif",
    "Article 3 — Description de l'existant et périmètre",
    "Article 4 — Exigences fonctionnelles (classe B)",
    "Article 5 — Architecture du système GTB",
    "Article 6 — Protocoles et interopérabilité",
    "Article 7 — Comptage et mesurage de l'énergie",
    "Article 8 — Cybersécurité (recommandations ANSSI)",
    "Article 9 — Protection des données (RGPD)",
    "Article 10 — Mise en service, commissioning et réception",
    "Article 11 — Maintenance, garanties et formation",
    "Article 12 — Documents à remettre et jugement des offres",
]
for s in sommaire:
    doc.add_paragraph(s, style="List Number")

doc.add_page_break()

# ---------- ARTICLE 1 ----------
article(1, "Objet du marché")
p = doc.add_paragraph()
p.add_run("Le présent Cahier des Clauses Techniques Particulières (CCTP) définit les exigences relatives "
          "à la fourniture, l'installation, le paramétrage, la mise en service et la réception d'un système "
          "de gestion technique du bâtiment (GTB), au sens du décret BACS, pour ")
todo(p, "[à compléter : désignation et adresse du bâtiment, maître d'ouvrage]")
p.add_run(".")
para("Le système installé devra atteindre a minima la classe B de la norme NF EN ISO 52120-1 (ex EN 15232), "
     "sur le périmètre défini à l'article 3. Le titulaire a une obligation de résultat sur l'atteinte de "
     "cette classe, qu'il devra démontrer à la réception.")

# ---------- ARTICLE 2 ----------
article(2, "Cadre réglementaire et normatif")
para("Les prestations sont réalisées dans le respect, notamment, des textes et normes suivants :")
for b in [
    "Décret n° 2020-887 du 20 juillet 2020 (« décret BACS »), transposant l'article 14 de la directive européenne EPBD 2018/844 ;",
    "Décret n° 2023-259 du 7 avril 2023 (abaissement du seuil de puissance CVC à 70 kW) ;",
    "Décret n° 2025-1343 du 26 décembre 2025 (report de la tranche 70–290 kW au 1er janvier 2030) ;",
    "Articles R. 175-1 à R. 175-6 du Code de la construction et de l'habitation ;",
    "Articles R. 241-26, R. 241-27 et R. 241-30 du Code de l'énergie (seuils et conditions de régulation) ;",
    "Norme NF EN ISO 52120-1 (mars 2022), qui remplace la norme EN 15232 ;",
    "Norme EN ISO 16484 pour le protocole BACnet ; normes EN 50090 et EN 13321 pour le protocole KNX.",
]:
    bullet(b)
p = doc.add_paragraph()
r = p.add_run("Classe visée : B. ")
r.bold = True
p.add_run("La classe C ne correspond qu'au minimum exigé pour certaines constructions neuves au titre "
          "d'autres réglementations ; elle est insuffisante pour le décret BACS.")

# ---------- ARTICLE 3 ----------
article(3, "Description de l'existant et périmètre")
p = doc.add_paragraph()
todo(p, "[à compléter] ")
p.add_run("Décrire le bâtiment et son patrimoine technique : surface, usage, nombre de niveaux, "
          "plages d'occupation.")
para("Préciser la puissance nominale utile cumulée des systèmes CVC (chauffage + climatisation), "
     "qui détermine l'assujettissement au décret BACS, ainsi que les lots concernés :")
for b in [
    "Production et distribution de chaleur (chaudières, PAC, sous-stations) ;",
    "Production et distribution de froid (groupes froid, PAC réversibles) ;",
    "Ventilation et traitement d'air (CTA, VMC, récupération) ;",
    "Eau chaude sanitaire (ECS) ;",
    "Éclairage des zones concernées ;",
    "Comptage et sous-comptage de l'énergie.",
]:
    bullet(b)
para("Décrire la régulation existante, les automates et protocoles déjà en place, et préciser si la "
     "prestation est une création, une extension ou une reprise d'installation.")

# ---------- ARTICLE 4 ----------
article(4, "Exigences fonctionnelles (classe B)")
para("La GTB devra assurer les trois finalités du décret BACS, déclinées en fonctions de classe B au "
     "sens de la norme NF EN ISO 52120-1 :")
for t, d in [
    ("A. Suivre, enregistrer, analyser et ajuster en continu",
     "Mesure et historisation continue des consommations par usage et par zone ; affichage des données "
     "d'exploitation ; ajustement automatique des consignes selon l'occupation réelle et les conditions extérieures."),
    ("B. Détecter les dérives et informer de l'efficacité énergétique",
     "Détection automatique des défauts et des écarts de performance des équipements ; génération d'alarmes "
     "hiérarchisées ; tableaux de bord et indicateurs permettant d'évaluer l'efficacité énergétique du bâtiment."),
    ("C. Garantir une GTB interopérable",
     "Communication avec les systèmes techniques du bâtiment via des protocoles ouverts et standardisés "
     "(article 6), sans dépendance à une solution propriétaire fermée."),
]:
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    para(d, space_after=8)
para("Fonctions d'automatisation attendues (classe B), par lot :")
for b in [
    "Régulation des émetteurs et générateurs CVC avec communication sur bus de terrain ;",
    "Programmation horaire centralisée et optimisation des relances selon l'occupation ;",
    "Régulation de la température ambiante par zone, avec consigne réduite en inoccupation ;",
    "Optimisation de la température de départ d'eau (loi d'eau) en fonction de l'extérieur ;",
    "Pilotage du débit d'air et de la vitesse des pompes et ventilateurs (variation de vitesse) ;",
    "Gestion centralisée de l'éclairage des zones concernées (présence et apport de lumière du jour) ;",
    "Comptage par usage et restitution d'indicateurs de consommation ;",
    "Fonctions de diagnostic et de détection automatique des défauts (FDD).",
]:
    bullet(b)

# ---------- ARTICLE 5 ----------
article(5, "Architecture du système GTB")
para("Le système sera organisé en trois niveaux, suivant l'architecture classique d'une GTB :")
for t, d in [
    ("Niveau terrain", "capteurs (température, CO2, présence, comptage) et actionneurs (vannes, registres, "
     "variateurs), raccordés sur bus de terrain."),
    ("Niveau automation", "automates programmables et contrôleurs assurant la régulation locale, le traitement "
     "des alarmes et l'historisation. Le fonctionnement de la régulation devra rester autonome en cas de perte de la supervision."),
    ("Niveau gestion", "poste de supervision (hyperviseur) centralisant la visualisation, les courbes de tendance, "
     "les alarmes, les rapports et le paramétrage, accessible localement et à distance de façon sécurisée."),
]:
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(t + " : ")
    r.bold = True
    p.add_run(d)
para("Les matériels seront standards, documentés et disponibles sur le marché. Aucune fonction essentielle "
     "ne devra dépendre d'un composant dont la maintenance serait réservée à un acteur unique.")

# ---------- ARTICLE 6 ----------
article(6, "Protocoles et interopérabilité")
para("Pour répondre à l'exigence d'interopérabilité du décret BACS, le système devra reposer sur des "
     "protocoles de communication ouverts et standardisés. Sont notamment acceptés :")
for b in [
    "BACnet (IP et MS/TP) — équipements et contrôleurs certifiés BTL ;",
    "Modbus (RTU et TCP) ;",
    "KNX ;",
    "LonWorks ;",
    "M-Bus (filaire ou radio) pour le comptage.",
]:
    bullet(b)
para("Les solutions entièrement propriétaires et fermées sont proscrites. Les passerelles éventuelles vers "
     "des équipements tiers seront documentées. Le maître d'ouvrage restera pleinement propriétaire de son "
     "installation, des codes sources de paramétrage et des accès, sans dépendance à un mainteneur unique.")

# ---------- ARTICLE 7 ----------
article(7, "Comptage et mesurage de l'énergie")
para("Le système assurera le comptage et le sous-comptage des consommations afin de permettre le suivi par "
     "usage exigé par le décret BACS :")
for b in [
    "Comptage par énergie (électricité, gaz, réseau de chaleur/froid, eau) ;",
    "Sous-comptage par usage significatif (chauffage, froid, ventilation, ECS, éclairage) ;",
    "Remontée des index et historisation des courbes de charge ;",
    "Compatibilité avec un report des données vers la plateforme OPERAT au titre du décret tertiaire, le cas échéant.",
]:
    bullet(b)

# ---------- ARTICLE 8 ----------
article(8, "Cybersécurité (recommandations ANSSI)")
para("L'installation s'appuiera sur les bonnes pratiques de sécurité des systèmes d'information, en s'inspirant "
     "des recommandations de l'ANSSI applicables aux systèmes industriels et à la GTB :")
for b in [
    "Segmentation du réseau GTB et cloisonnement vis-à-vis du réseau bureautique ;",
    "Gestion des comptes et des droits par profils (consultation, modification, acquittement des alarmes, administration) ;",
    "Authentification forte et chiffrement des accès distants ;",
    "Politique de mises à jour de sécurité et de sauvegarde des configurations ;",
    "Journalisation des accès et des actions sur le système.",
]:
    bullet(b)

# ---------- ARTICLE 9 ----------
article(9, "Protection des données (RGPD)")
para("Si le système traite des données à caractère personnel (par exemple via la détection de présence ou des "
     "comptages individualisés), le titulaire devra respecter le RGPD : minimisation des données, finalité "
     "limitée à la gestion technique et énergétique, durée de conservation maîtrisée et information des personnes "
     "concernées. Aucune donnée ne sera exploitée à des fins étrangères à l'exploitation du bâtiment.")

# ---------- ARTICLE 10 ----------
article(10, "Mise en service, commissioning et réception")
for b in [
    "Programme de mise en service détaillant les tests par lot et par point ;",
    "Commissioning fonctionnel : vérification du fonctionnement réel des séquences de régulation et d'optimisation ;",
    "Démonstration de l'atteinte de la classe B au sens de la norme NF EN ISO 52120-1, fonction par fonction ;",
    "Vérification des alarmes, des historiques et des indicateurs énergétiques ;",
    "Levée des réserves avant prononcé de la réception.",
]:
    bullet(b)

# ---------- ARTICLE 11 ----------
article(11, "Maintenance, garanties et formation")
for b in [
    "Garantie de parfait achèvement et garantie des matériels ;",
    "Contrat de maintenance préventive et curative (à préciser : durée, délais d'intervention) ;",
    "Formation des exploitants à la supervision et au paramétrage de premier niveau ;",
    "Remise des accès, mots de passe et codes de paramétrage au maître d'ouvrage, garantissant son autonomie "
    "et la possibilité de changer de mainteneur.",
]:
    bullet(b)

# ---------- ARTICLE 12 ----------
article(12, "Documents à remettre et jugement des offres")
para("Le titulaire remettra notamment : la liste des points (entrées/sorties), les analyses fonctionnelles, les "
     "schémas d'architecture et de réseau, les synoptiques de supervision et le dossier des ouvrages exécutés (DOE).")
p = doc.add_paragraph()
todo(p, "[à compléter] ")
p.add_run("Critères de jugement suggérés :")
for b in [
    "Valeur technique (architecture, ouverture des protocoles, démonstration de l'atteinte de la classe B) ;",
    "Prix des prestations ;",
    "Qualité du mémoire technique et des références équivalentes ;",
    "Délais et conditions de maintenance.",
]:
    bullet(b)

# ---------- PIED ----------
doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run("Modèle élaboré par NeoGTB à partir de vrais CCTP de consultations publiques traitant du "
                  "décret BACS et de la norme NF EN ISO 52120-1. Trame neutre, sans lien fabricant. "
                  "Pour une version calée sur votre bâtiment : neogtb.fr/amo-gtb-gtc.")
fr.italic = True
fr.font.size = Pt(9)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print("OK ->", OUT)
